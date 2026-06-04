"""Ingest a company document into the local RAG store.

Usage:  python -m scripts.ingest_doc "C:\\path\\to\\policy.docx" [--name "Travel Policy"]

Supports .docx and .txt. Company policies are stored as acl_scope='general'
(readable by all staff); the AI's search_company_policy tool retrieves them.
"""
import sys
from pathlib import Path

from app.core.db import Base, SessionLocal, engine
from app.modules import ai  # noqa: F401  register rag_chunks
from app.modules.ai import rag


def docx_text(path: str) -> str:
    import docx

    d = docx.Document(path)
    parts: list[str] = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return docx_text(path)
    if ext in (".txt", ".md"):
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    raise SystemExit(f"unsupported file type: {ext} (use .docx or .txt)")


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit("usage: python -m scripts.ingest_doc <file> [--name NAME]")
    path = sys.argv[1]
    name = "Travel & Expense Policy"
    if "--name" in sys.argv:
        name = sys.argv[sys.argv.index("--name") + 1]

    text = extract(path)
    print(f"Extracted {len(text)} chars from {Path(path).name}")
    Base.metadata.create_all(engine)
    with SessionLocal() as s:
        n = rag.ingest(s, source=name, text=text, acl_scope="general")
        s.commit()
    print(f"Ingested '{name}' as {n} chunks (embedded with bge-m3).")


if __name__ == "__main__":
    main()
